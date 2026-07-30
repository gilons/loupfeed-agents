"""Tools: ``graph_api``, ``graph_meeting_transcript``, ``graph_file_content``.
Read-only Microsoft Graph.

The pm agent is installed into Teams teams/chats/meetings, where resource-
specific consent (RSC) grants its Entra app read access to that resource's
messages, members, and meeting artifacts — plus whatever tenant application
permissions were consented in Entra (directory, SharePoint sites). GET-only by
design, mirroring ``github_api``: writes stay with the Bot Framework reply
path and the planning-system connectors.
"""

from __future__ import annotations

import base64
import io
import json
import logging
from typing import Any
from urllib.parse import quote

import requests

from ..utils.msgraph import GRAPH_BASE, get_graph_app_token

logger = logging.getLogger(__name__)

_MAX_RESPONSE_CHARS = 60_000
_MAX_TRANSCRIPT_CHARS = 120_000
_MAX_FILE_BYTES = 15 * 1024 * 1024
_TIMEOUT = 30

_TEXT_EXTENSIONS = {
    "txt", "md", "csv", "tsv", "json", "yaml", "yml", "xml", "html", "htm", "log",
    "py", "js", "ts", "java", "go", "rb", "sh", "sql", "toml", "ini", "cfg",
}  # fmt: skip

# Read-oriented Graph roots the pm agent has any business touching.
_ALLOWED_ROOTS = ("chats", "teams", "groups", "users", "sites", "drives")


def _fail(reason: str, *, detail: str = "", where: str = "") -> dict[str, Any]:
    """Return a plain-language failure, keeping the diagnostics in the server log.

    Whatever lands in ``reason`` can end up quoted verbatim to whoever asked, so
    it stays free of status codes, endpoints, permission names and tracebacks —
    those read as noise to the asker and as an invitation to go chase the wrong
    fix. The detail we actually need is logged instead.
    """
    if detail:
        logger.warning("graph tool failed%s: %s", f" [{where}]" if where else "", detail)
    return {"ok": False, "reason": reason}


def _headers() -> dict[str, str] | None:
    token = get_graph_app_token()
    if not token:
        return None
    return {"Authorization": f"Bearer {token}"}


def _allowed(path: str) -> bool:
    root = path.lstrip("/").split("/", 1)[0].split("?", 1)[0].lower()
    return root in _ALLOWED_ROOTS


def _get(path: str, params: dict[str, Any] | None = None) -> requests.Response:
    headers = _headers()
    if headers is None:
        raise RuntimeError("Microsoft Graph credentials not configured")
    return requests.get(
        f"{GRAPH_BASE}{path}", headers=headers, params=params or {}, timeout=_TIMEOUT
    )


def graph_api(path: str, params: dict[str, Any] | None = None) -> dict[str, Any]:
    """Call a read-only (GET) Microsoft Graph v1.0 endpoint as the loupfeed Teams app.

    Access is scoped by where the app is installed (Teams RSC) plus tenant
    read permissions — use it to READ Microsoft 365 context instead of asking
    people to paste it. Useful endpoints:

    - ``/chats/{chat-id}/messages`` — recent messages in a group/meeting chat
    - ``/chats/{chat-id}/members`` — who is in the chat
    - ``/chats/{chat-id}`` — chat properties (a meeting chat carries ``onlineMeetingInfo``)
    - ``/teams/{team-group-id}/channels`` — a team's channels
    - ``/teams/{team-group-id}/channels/{channel-id}/messages`` — channel messages
    - ``/teams/{team-group-id}/members`` — team roster
    - ``/groups/{team-group-id}/drive/root/children`` — files in the team's library
    - ``/users`` with ``$filter=startswith(displayName,'...')`` — resolve people
    - ``/users/{id-or-upn}`` / ``/users/{id}/manager`` — profile and org chart
    - ``/sites/{hostname}:/{site-path}`` — SharePoint sites

    Only GET is possible. Chat/channel ids come from the Teams context of the
    current conversation (see the system prompt) or from other Graph calls.

    Args:
        path: The API path, starting with ``/`` (e.g. ``/chats/{id}/messages``).
        params: Optional query parameters (``$top``, ``$filter``, ``$select``, ...).

    Returns:
        ``{"status": int, "body": <parsed JSON or text>}``. 403 usually means
        the app is not installed in that team/chat (no RSC grant there) or the
        tenant permission is missing — stop rather than retrying, and tell the
        asker plainly that you can't reach it. Never quote status codes, Graph
        paths, permission names or error bodies into a reply: they are logged
        server-side for us, and they send people chasing the wrong fix.
    """
    if not isinstance(path, str) or not path.startswith("/"):
        return {"status": 0, "body": "path must start with '/' (e.g. /chats/{id}/messages)"}
    if not _allowed(path):
        return {
            "status": 0,
            "body": f"path root not allowed; must be one of: {', '.join(_ALLOWED_ROOTS)}",
        }
    try:
        resp = _get(path, params)
    except RuntimeError as exc:
        logger.warning("graph_api %s failed: %s: %s", path, type(exc).__name__, exc)
        return {"status": 0, "body": str(exc)}
    except requests.RequestException as exc:
        logger.warning("graph_api %s failed: %s: %s", path, type(exc).__name__, exc)
        return {"status": 0, "body": f"request failed: {exc}"}

    if resp.status_code >= 400:
        logger.warning("graph_api %s -> %s: %s", path, resp.status_code, resp.text[:500])

    if resp.status_code == 204 or not resp.content:
        return {"status": resp.status_code, "body": ""}
    try:
        body: Any = resp.json()
        text = json.dumps(body)
    except ValueError:
        text = resp.text
        body = text
    if len(text) > _MAX_RESPONSE_CHARS:
        return {"status": resp.status_code, "body": text[:_MAX_RESPONSE_CHARS], "truncated": True}
    return {"status": resp.status_code, "body": body}


def _extract_file_text(name: str, mime: str, data: bytes) -> tuple[str, str]:
    """Return ``(kind, text)`` for a downloaded file, or raise ValueError."""
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if ext == "pdf" or mime == "application/pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "pdf", "\n\n".join(pages).strip()
    if ext == "docx" or mime.endswith("wordprocessingml.document"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "docx", "\n".join(parts)
    if mime.startswith("text/") or ext in _TEXT_EXTENSIONS:
        return "text", data.decode("utf-8", errors="replace")
    raise ValueError(
        f"unsupported file type ({name or mime or 'unknown'}); "
        "readable types: pdf, docx, and plain-text formats"
    )


def graph_find_recording(
    team_id: str, channel_name: str, on_or_before: str = "", subject_contains: str = ""
) -> dict[str, Any]:
    """Find a channel meeting's recording and return a link that can be transcribed.

    Teams puts a channel meeting's recording in that channel's SharePoint folder
    (``<channel>/Recordings``), which this app can read. Use this when someone asks
    what happened in a call held in a channel — a standup, a review — and then pass
    the returned ``audio_url`` to ``transcribe_recording``.

    Recordings of meetings that were NOT held in a channel live in the organiser's
    personal OneDrive and are not reachable this way.

    Args:
        team_id: The team's group id (in the Teams context of this conversation).
        channel_name: The channel's display name, e.g. ``Dev team``.
        on_or_before: Optional ``YYYY-MM-DD``; returns the newest recording created
            on or before that date. Omit for the most recent one.
        subject_contains: Optional case-insensitive filter on the file name, e.g.
            ``Daily Standup``, for channels hosting several meeting series.

    Returns:
        ``{"ok": True, "name": str, "created": str, "size": int, "audio_url": str,
        "candidates": [names...]}`` or ``{"ok": False, "reason": str}``. The
        ``audio_url`` is short-lived — use it immediately and never log it.
    """
    if not team_id or not channel_name:
        return _fail("I need to know which team and channel the meeting was in.")
    folder = f"{channel_name}/Recordings"
    try:
        resp = _get(f"/groups/{team_id}/drive/root:/{quote(folder)}:/children")
        if resp.status_code != 200:
            return _fail(
                f"I couldn't find a recordings folder for the {channel_name} channel.",
                detail=f"list {folder} -> {resp.status_code}: {resp.text[:400]}",
                where="find_recording.list",
            )
        items = [
            i for i in (resp.json().get("value") or []) if str(i.get("name", "")).endswith(".mp4")
        ]
        if subject_contains:
            needle = subject_contains.lower()
            items = [i for i in items if needle in str(i.get("name", "")).lower()]
        if on_or_before:
            items = [i for i in items if str(i.get("createdDateTime", ""))[:10] <= on_or_before]
        if not items:
            return _fail("I couldn't find a recording matching that in this channel.")
        items.sort(key=lambda i: str(i.get("createdDateTime") or ""), reverse=True)
        newest = items[0]
        url = newest.get("@microsoft.graph.downloadUrl")
        if not url:
            return _fail(
                "I found the recording but can't get a link to its audio.",
                detail=f"no downloadUrl on item {newest.get('id')}",
                where="find_recording.url",
            )
        return {
            "ok": True,
            "name": str(newest.get("name") or ""),
            "created": str(newest.get("createdDateTime") or ""),
            "size": int(newest.get("size") or 0),
            "audio_url": url,
            "candidates": [str(i.get("name") or "") for i in items[:5]],
        }
    except RuntimeError as exc:
        return _fail(
            "I'm not able to reach Teams right now.",
            detail=f"{type(exc).__name__}: {exc}",
            where="find_recording",
        )
    except requests.RequestException as exc:
        return _fail(
            "I couldn't reach Teams to look for the recording.",
            detail=f"{type(exc).__name__}: {exc}",
            where="find_recording",
        )


def graph_file_content(url: str = "", drive_id: str = "", item_id: str = "") -> dict[str, Any]:
    """Download a file shared in Teams/SharePoint and return its readable text.

    Use this to actually READ a shared document instead of only listing it.
    Handles PDF, DOCX, and plain-text formats (md, csv, json, code, ...).

    Pass ONE of:
    - ``url`` — the file's web link, e.g. an attachment's ``contentUrl`` from a
      channel/chat message or a SharePoint link someone pasted.
    - ``drive_id`` + ``item_id`` — from a Graph drive listing
      (``/groups/{team}/drive/root/children``).

    Files in a team's SharePoint library are readable wherever the app is
    installed. Files shared in private/group chats live in the sharer's
    personal OneDrive and may be denied (403) unless the tenant granted
    ``Files.Read.All`` — report that instead of retrying.

    Returns:
        ``{"ok": True, "name": str, "kind": "pdf|docx|text", "size": int,
        "text": str}`` (``truncated: True`` when capped) or
        ``{"ok": False, "reason": str}``.
    """
    if not (url or (drive_id and item_id)):
        return {"ok": False, "reason": "pass url, or drive_id and item_id"}
    try:
        if not (drive_id and item_id):
            share_id = "u!" + base64.urlsafe_b64encode(url.encode()).decode().rstrip("=")
            item_resp = _get(f"/shares/{share_id}/driveItem")
            if item_resp.status_code != 200:
                return {
                    "ok": False,
                    "reason": (
                        f"could not resolve the file link ({item_resp.status_code}): "
                        f"{item_resp.text[:300]}"
                    ),
                }
            item = item_resp.json()
            drive_id = ((item.get("parentReference") or {}).get("driveId")) or ""
            item_id = item.get("id") or ""
            if not drive_id or not item_id:
                return {"ok": False, "reason": "link resolved but not to a drive file"}
        else:
            item_resp = _get(f"/drives/{drive_id}/items/{item_id}")
            if item_resp.status_code != 200:
                return {
                    "ok": False,
                    "reason": (
                        f"could not read file metadata ({item_resp.status_code}): "
                        f"{item_resp.text[:300]}"
                    ),
                }
            item = item_resp.json()

        name = str(item.get("name") or "")
        mime = str(((item.get("file") or {}).get("mimeType")) or "")
        size = int(item.get("size") or 0)
        if size > _MAX_FILE_BYTES:
            return {
                "ok": False,
                "reason": f"file too large ({size} bytes; limit {_MAX_FILE_BYTES})",
            }

        content_resp = _get(f"/drives/{drive_id}/items/{item_id}/content")
        if content_resp.status_code != 200:
            return {
                "ok": False,
                "reason": (
                    f"could not download the file ({content_resp.status_code}): "
                    f"{content_resp.text[:300]}"
                ),
            }
        data = content_resp.content
        if len(data) > _MAX_FILE_BYTES:
            return {
                "ok": False,
                "reason": f"file too large ({len(data)} bytes; limit {_MAX_FILE_BYTES})",
            }

        try:
            kind, text = _extract_file_text(name, mime, data)
        except ValueError as exc:
            return {"ok": False, "reason": str(exc)}
        except Exception as exc:  # noqa: BLE001 — corrupt/odd files must not crash the run
            return {"ok": False, "reason": f"could not parse {name or 'file'}: {exc}"}

        truncated = len(text) > _MAX_RESPONSE_CHARS
        return {
            "ok": True,
            "name": name,
            "kind": kind,
            "size": size or len(data),
            "text": text[:_MAX_RESPONSE_CHARS],
            **({"truncated": True} if truncated else {}),
        }
    except RuntimeError as exc:
        return {"ok": False, "reason": str(exc)}
    except requests.RequestException as exc:
        return {"ok": False, "reason": f"request failed: {exc}"}


def graph_meeting_transcript(chat_id: str) -> dict[str, Any]:
    """Fetch the transcript of the Teams meeting behind a meeting chat.

    Use this right after a call, when mentioned in a meeting chat: pass the
    conversation's chat id (from the Teams context in the system prompt). It
    resolves the chat → its online meeting → the newest transcript and returns
    the transcript text (WebVTT: timestamped, speaker-attributed lines).

    Only works for scheduled meetings with transcription turned on; returns a
    clear reason otherwise (no meeting behind the chat, transcription off, or
    access not granted).

    Args:
        chat_id: The meeting chat id (looks like ``19:meeting_...@thread.v2``).

    Returns:
        ``{"ok": bool, "subject": str, "organizer_id": str, "transcript": str}``
        on success; ``{"ok": False, "reason": ...}`` on failure.
    """
    if not isinstance(chat_id, str) or not chat_id:
        return {"ok": False, "reason": "chat_id is required"}
    try:
        chat_resp = _get(f"/chats/{chat_id}")
        if chat_resp.status_code in (403, 404):
            return _fail(
                "I can't reach this meeting's transcript — I don't have access to this "
                "conversation. Someone will need to add me to it. If you paste the "
                "transcript or a summary, I can work from that instead.",
                detail=f"GET /chats/{chat_id} -> {chat_resp.status_code}: {chat_resp.text[:500]}",
                where="meeting_transcript.chat",
            )
        if chat_resp.status_code != 200:
            return _fail(
                "I couldn't read this meeting's conversation, so I can't get to the "
                "transcript. Paste it or summarise the call and I'll take it from there.",
                detail=f"GET /chats/{chat_id} -> {chat_resp.status_code}: {chat_resp.text[:500]}",
                where="meeting_transcript.chat",
            )
        chat = chat_resp.json()
        meeting_info = chat.get("onlineMeetingInfo") or {}
        join_url = meeting_info.get("joinWebUrl")
        organizer_id = ((meeting_info.get("organizer") or {}).get("id")) or ""
        if not join_url or not organizer_id:
            return {"ok": False, "reason": "this chat has no online meeting associated with it"}

        escaped = join_url.replace("'", "''")
        meetings_resp = _get(
            f"/users/{organizer_id}/onlineMeetings",
            params={"$filter": f"JoinWebUrl eq '{escaped}'"},
        )
        if meetings_resp.status_code != 200:
            return _fail(
                "I couldn't look up the meeting behind this conversation, so the "
                "transcript is out of reach. Paste it and I'll carry on.",
                detail=(
                    f"onlineMeetings lookup -> {meetings_resp.status_code}: "
                    f"{meetings_resp.text[:500]}"
                ),
                where="meeting_transcript.resolve",
            )
        meetings = meetings_resp.json().get("value") or []
        if not meetings:
            return _fail("I couldn't find the meeting behind this conversation.")
        meeting = meetings[0]

        transcripts_resp = _get(f"/users/{organizer_id}/onlineMeetings/{meeting['id']}/transcripts")
        if transcripts_resp.status_code != 200:
            return _fail(
                "I can't get to this meeting's transcript. Paste it or summarise the "
                "call and I'll take it from there.",
                detail=(
                    f"list transcripts -> {transcripts_resp.status_code}: "
                    f"{transcripts_resp.text[:500]}"
                ),
                where="meeting_transcript.list",
            )
        transcripts = transcripts_resp.json().get("value") or []
        if not transcripts:
            return _fail(
                "There's no transcript for this meeting — transcription may have been off."
            )
        transcripts.sort(key=lambda t: str(t.get("createdDateTime") or ""), reverse=True)
        newest = transcripts[0]

        content_resp = _get(
            f"/users/{organizer_id}/onlineMeetings/{meeting['id']}"
            f"/transcripts/{newest['id']}/content",
            params={"$format": "text/vtt"},
        )
        if content_resp.status_code != 200:
            return _fail(
                "I found the transcript but couldn't download it. Paste it and I'll carry on.",
                detail=(
                    f"transcript content -> {content_resp.status_code}: {content_resp.text[:500]}"
                ),
                where="meeting_transcript.content",
            )
        text = content_resp.text
        truncated = len(text) > _MAX_TRANSCRIPT_CHARS
        return {
            "ok": True,
            "subject": str(meeting.get("subject") or chat.get("topic") or ""),
            "organizer_id": organizer_id,
            "created": str(newest.get("createdDateTime") or ""),
            "transcript": text[:_MAX_TRANSCRIPT_CHARS],
            **({"truncated": True} if truncated else {}),
        }
    except RuntimeError as exc:
        return _fail(
            "I'm not able to reach Teams right now.",
            detail=f"{type(exc).__name__}: {exc}",
            where="meeting_transcript",
        )
    except requests.RequestException as exc:
        return _fail(
            "I couldn't reach Teams to fetch the transcript. Worth retrying in a moment.",
            detail=f"{type(exc).__name__}: {exc}",
            where="meeting_transcript",
        )
